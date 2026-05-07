"""
Document processing service.
Extracts clean plain text from PDF, DOCX, and TXT files.
"""
import io
import re
from pathlib import Path
from fastapi import UploadFile, HTTPException
from backend.utils.logger import setup_logger

logger = setup_logger(__name__)

SUPPORTED_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/msword": ".doc",
    "text/plain": ".txt",
    "text/markdown": ".md",
}


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove control characters."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except ImportError:
        logger.warning("pdfplumber not installed, trying fallback PDF extraction")
        return _extract_pdf_fallback(content)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise HTTPException(status_code=422, detail=f"Failed to parse PDF: {str(e)}")


def _extract_pdf_fallback(content: bytes) -> str:
    """Fallback PDF extraction using PyMuPDF if available."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        return "\n\n".join(page.get_text() for page in doc)
    except ImportError:
        raise HTTPException(
            status_code=422,
            detail="PDF parsing requires pdfplumber. Run: pip install pdfplumber"
        )


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX preserving paragraph structure."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name if para.style else ""
                # Add heading markers for context
                if "Heading" in style:
                    parts.append(f"\n### {para.text.strip()}\n")
                else:
                    parts.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n\n".join(parts)
    except ImportError:
        raise HTTPException(
            status_code=422,
            detail="DOCX parsing requires python-docx. Run: pip install python-docx"
        )
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise HTTPException(status_code=422, detail=f"Failed to parse DOCX: {str(e)}")


def _extract_txt(content: bytes) -> str:
    """Decode plain text with encoding detection fallback."""
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


async def extract_text_from_upload(file: UploadFile) -> str:
    """
    Main entry point: reads an uploaded file and returns clean extracted text.
    Raises HTTPException for unsupported types or extraction failures.
    """
    content = await file.read()

    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    filename = file.filename or ""
    content_type = file.content_type or ""
    suffix = Path(filename).suffix.lower()

    logger.info(f"Processing document: {filename} ({content_type}, {len(content)} bytes)")

    # Determine type by extension first, then content-type
    if suffix == ".pdf" or "pdf" in content_type:
        raw_text = _extract_pdf(content)
    elif suffix in (".docx", ".doc") or "wordprocessingml" in content_type or "msword" in content_type:
        raw_text = _extract_docx(content)
    elif suffix in (".txt", ".md") or content_type.startswith("text/"):
        raw_text = _extract_txt(content)
    else:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type: '{suffix}'. Supported: PDF, DOCX, TXT, MD"
        )

    cleaned = _clean_text(raw_text)

    if len(cleaned) < 20:
        raise HTTPException(
            status_code=422,
            detail="Could not extract meaningful text from the document. "
                   "Ensure the file contains readable text (not scanned images)."
        )

    logger.info(f"Extracted {len(cleaned)} characters from {filename}")
    return cleaned
